# agents/cam_generator.py
# AGENT 7 — Synthesizes all outputs into a professional CAM document

import sys, os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from utils.llm_client import call_llm
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

SYNTHESIS_PROMPT = """
You are SENTINEL's CAM Generation Agent. Synthesize all agent outputs into 
a professional, bank-grade Credit Appraisal Memorandum.

Write like a senior banker with 20 years experience — precise, formal, 
evidence-based. Every claim must reference specific data.

Produce the CAM in this EXACT structure:

SECTION 1: EXECUTIVE SUMMARY
[3 sentences: Decision | Amount | Rate | Primary Reason]

SECTION 2: BORROWER PROFILE
2.1 Company Overview [incorporated, CIN, business, 4 lines]
2.2 Promoter Profile [background, track record, 4 lines]
2.3 Shareholding Pattern [promoter %, public %, pledged %]

SECTION 3: FACILITY REQUESTED
Type | Amount | Purpose | Tenure | Security proposed | Repayment schedule

SECTION 4: FIVE Cs ANALYSIS

4A. CHARACTER [Score: __/20]
[3 specific evidence points supporting your character assessment]

4B. CAPACITY [Score: __/25]
[3-year revenue/EBITDA trend | DSCR | Cash flow quality]

4C. CAPITAL [Score: __/20]
[D/E ratio | Net worth adequacy | Promoter's own stake]

4D. COLLATERAL [Score: __/15]
[Type | Value | Coverage ratio | Existing charges]

4E. CONDITIONS [Score: __/20]
[Sector outlook | Regulatory environment | Macro risks]

SECTION 5: FINANCIAL ANALYSIS
[3-year income statement highlights | Key ratios vs benchmarks]

SECTION 6: DATA INTEGRITY ASSESSMENT
[Cross-verification results | Fraud pattern scan summary | Confidence level]

SECTION 7: EXTERNAL INTELLIGENCE SUMMARY
[Key research findings | Promoter network risk | Sector risk]

SECTION 8: PRIMARY DUE DILIGENCE
[Site visit observations | How they changed the score]

SECTION 9: STRESS TEST SUMMARY
[Table of 4 scenarios | Survival rating | Covenants triggered]

SECTION 10: RISK FACTORS & MITIGANTS
[Table: Risk | Severity | Mitigation]

SECTION 11: THE COMMITTEE DEBATE
Bull Agent's top 3 arguments: [with evidence]
Bear Agent's top 3 arguments: [with evidence]
Chairman's resolution: [why one side won]

SECTION 12: CREDIT SCORECARD
[Full pillar-by-pillar breakdown]

SECTION 13: FINAL CREDIT DECISION
[Decision | Amount | Rate | Tenure | Security | Conditions | Monitoring triggers]
"""


def synthesize_cam_text(all_outputs: dict) -> str:
    """
    Uses LLM to synthesize all agent outputs into coherent CAM text.
    """
    loan = all_outputs.get('loan_details', {})
    
    user_message = f"""
    Synthesize all the following agent outputs into a complete, professional 
    Credit Appraisal Memorandum for {loan.get('company_name', 'the company')}.
    
    Write in formal banking language. Every section must have specific data points.
    Do NOT write generic statements — reference actual numbers and findings.
    
    DOCUMENT PARSER OUTPUT:
    {all_outputs.get('parser', 'Not available')[:2500]}
    
    RESEARCH INTELLIGENCE:
    {all_outputs.get('research', 'Not available')[:2500]}
    
    FRAUD DETECTION:
    {all_outputs.get('fraud', 'Not available')[:2500]}
    
    BULL AGENT BRIEF:
    {all_outputs.get('bull', 'Not available')[:2500]}
    
    BEAR AGENT BRIEF:
    {all_outputs.get('bear', 'Not available')[:2500]}
    
    CHAIRMAN'S VERDICT (FINAL DECISION):
    {all_outputs.get('chairman', 'Not available')[:2500]}
    
    STRESS TEST RESULTS:
    {all_outputs.get('stress_test', 'Not available')[:2500]}
    
    PRIMARY DUE DILIGENCE NOTES:
    {all_outputs.get('primary_notes', 'No site visit notes provided.')[:2500]}
    """
    
    return call_llm("cam_generator", SYNTHESIS_PROMPT, user_message)


def create_cam_word_document(cam_text: str, loan_details: dict, 
                              output_path: str = None) -> str:
    """
    Creates a professionally formatted Word document from the CAM text.
    
    Returns: Path to the generated .docx file
    """
    doc = Document()
    
    # ── Page Setup ───────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Helper functions ─────────────────────────────────────────
    def add_heading(text, level=1, color=(0, 51, 102)):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
        return p

    def add_paragraph(text, bold=False, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    def add_horizontal_line():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '003366')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── TITLE PAGE ───────────────────────────────────────────────
    title = doc.add_heading('CREDIT APPRAISAL MEMORANDUM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.font.size = Pt(20)

    doc.add_paragraph()
    
    # Metadata table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    meta = [
        ('Prepared by:', 'SENTINEL Credit Intelligence System v1.0'),
        ('Company:', loan_details.get('company_name', 'N/A')),
        ('Loan Amount:', f"₹{loan_details.get('loan_amount', 'N/A')} crore"),
        ('Date:', datetime.now().strftime('%d %B %Y')),
        ('Classification:', 'STRICTLY CONFIDENTIAL'),
    ]
    for i, (key, val) in enumerate(meta):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ── DISCLAIMER BOX ───────────────────────────────────────────
    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        "⚠️  IMPORTANT: This memorandum is AI-generated by SENTINEL and constitutes "
        "a recommendation only. Final sanction authority rests with the authorized "
        "Credit Officer. All primary document verification mandatory before disbursement."
    ).bold = True
    disclaimer.style.font.size = Pt(9)

    add_horizontal_line()
    doc.add_paragraph()

    # ── CAM CONTENT ──────────────────────────────────────────────
    # Parse the LLM-generated text and format it section by section
    lines = cam_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        # Detect section headers
        if line.startswith('SECTION') and ':' in line:
            add_heading(line, level=1)
            add_horizontal_line()
        elif line.startswith(('4A.', '4B.', '4C.', '4D.', '4E.')):
            add_heading(line, level=2, color=(0, 102, 51))
        elif line.startswith(('2.', '3.', '5.', '6.', '7.', '8.', '9.', '10.')):
            add_heading(line, level=2, color=(51, 51, 102))
        elif line.startswith('===') or line.startswith('---'):
            add_horizontal_line()
        elif line.startswith('│') or line.startswith('┌') or line.startswith('└'):
            # Table-like content — use monospace
            p = doc.add_paragraph(line)
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)
        elif line.startswith('🔴') or line.startswith('🟠') or line.startswith('🟢'):
            # Risk indicators — make bold
            add_paragraph(line, bold=True)
        else:
            doc.add_paragraph(line)

    # ── FOOTER ───────────────────────────────────────────────────
    doc.add_page_break()
    footer_text = doc.add_paragraph()
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_text.add_run(
        "SENTINEL v1.0 | Adversarial Credit Intelligence Platform\n"
        '"We don\'t just read documents. We interrogate them."'
    )
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.size = Pt(9)

    # ── SAVE ─────────────────────────────────────────────────────
    if not output_path:
        company = loan_details.get('company_name', 'company').replace(' ', '_').replace('(', '').replace(')', '')
        date_str = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = f"output/CAM_{company}_{date_str}.docx"
    
    # Ensure output directory exists
    output_dir = os.path.dirname(output_path) or "output"
    try:
        os.makedirs(output_dir, exist_ok=True)
    except Exception as e:
        print(f"⚠️  Warning: Could not create directory {output_dir}: {e}")
        output_dir = "output"
        os.makedirs(output_dir, exist_ok=True)
    
    # Try to save the file with retry logic
    save_successful = False
    max_retries = 3
    
    for attempt in range(max_retries):
        try:
            doc.save(output_path)
            print(f"✅ CAM Document saved: {output_path}")
            save_successful = True
            break
        except PermissionError as e:
            if attempt < max_retries - 1:
                # File might be open in Word, try with a different name
                print(f"⚠️  Attempt {attempt + 1}/{max_retries}: File locked, retrying with different name...")
                base_name = os.path.basename(output_path).replace('.docx', '')
                output_path = os.path.join(
                    output_dir, 
                    f"{base_name}_v{attempt + 2}.docx"
                )
            else:
                print(f"❌ Permission denied: {output_path}")
                print(f"   Hint: Close any open CAM documents in Word and try again")
                raise
        except Exception as e:
            print(f"❌ Error saving CAM document: {str(e)}")
            raise
    
    if not save_successful:
        raise RuntimeError(f"Could not save CAM document after {max_retries} attempts")
    
    return output_path


def run_cam_generator(all_outputs: dict) -> tuple:
    """
    Master function — synthesizes all agent outputs and creates CAM document.
    
    Args:
        all_outputs: dict containing all agent outputs + loan_details
    
    Returns:
        Tuple of (cam_text: str, doc_path: str)
    """
    print("📄 Running CAM Generator Agent...")
    
    # Step 1: Synthesize all outputs into structured CAM text
    cam_text = synthesize_cam_text(all_outputs)
    
    # Step 2: Create formatted Word document
    # Pass None so create_cam_word_document generates a unique filename with timestamp
    doc_path = create_cam_word_document(
        cam_text=cam_text,
        loan_details=all_outputs.get('loan_details', {}),
        output_path=None  # Let the function generate a unique filename
    )
    
    print("✅ CAM Generation Complete")
    return cam_text, doc_path